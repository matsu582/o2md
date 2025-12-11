#!/usr/bin/env python3
"""
Word to Markdown Converter
シンプルで確実なWord文書のMarkdown変換ツール

特徴:
- 本文、表、図の位置をWord文書通りに再現
- 表、箇条書き、段落番号、見出しに対応
- 目次の自動生成
- Word図形と画像の組み合わせを図として変換
- 単純な画像はそのまま表示
- 見出し参照のリンク対応
"""

import os
import sys
import re
import tempfile
import subprocess
import math
import shutil
import zipfile
import urllib.parse
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
from PIL import Image
import io

from utils import get_libreoffice_path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    raise ImportError(
        "python-docxライブラリが必要です: pip install python-docx または uv sync を実行してください"
    ) from e

try:
    from PIL import Image
except ImportError as e:
    raise ImportError(
        "Pillowライブラリが必要です: pip install pillow または uv sync を実行してください"
    ) from e

try:
    import fitz
except ImportError as e:
    raise ImportError(
        "PyMuPDFライブラリが必要です: pip install PyMuPDF または uv sync を実行してください"
    ) from e

# 設定
LIBREOFFICE_PATH = get_libreoffice_path()


# グローバルverboseフラグ
_VERBOSE = False

def set_verbose(verbose: bool):
    """verboseモードを設定"""
    global _VERBOSE
    _VERBOSE = verbose

def is_verbose() -> bool:
    """verboseモードかどうかを返す"""
    return _VERBOSE

def debug_print(*args, **kwargs):
    """verboseモード時のみ出力するデバッグ用print"""
    if _VERBOSE:
        print(*args, **kwargs)

class WordToMarkdownConverter:
    def __init__(self, word_file_path: str, use_heading_text=False, output_dir=None, shape_metadata=False, output_format='png'):
        """Word文書をMarkdownに変換するコンバータ
        
        Args:
            word_file_path: 変換するWordファイルのパス
            use_heading_text: 章番号の代わりに見出しテキストを使用するか
            output_dir: 出力ディレクトリ（省略時はデフォルト）
            shape_metadata: 図形メタデータ出力フラグ
            output_format: 出力画像形式 ('png' または 'svg')
        """
        self.word_file = word_file_path
        self.doc = Document(word_file_path)
        self.base_name = Path(word_file_path).stem
        
        # 出力ディレクトリの設定
        if output_dir:
            self.output_dir = output_dir
        else:
            self.output_dir = os.path.join(os.getcwd(), "output")
        
        self.images_dir = os.path.join(self.output_dir, "images")
        
        # ディレクトリ作成
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.images_dir, exist_ok=True)
        
        self.markdown_lines = []
        self.headings = []  # 見出し一覧
        self.headings_map = {}  # 章番号とアンカーのマッピング
        self.heading_titles_map = {}  # 章番号と見出しタイトルのマッピング
        self.use_heading_text = use_heading_text  # 章番号の代わりに見出しテキストを使用するオプション
        self.processed_images = {}  # ハッシュベース重複検出用辞書
        self.referenced_images = set()  # 実際に文書内で参照されている画像のrId
        self.vector_image_counter = 0  # ベクター画像専用カウンター
        self.regular_image_counter = 0  # 通常画像専用カウンター
        self.shape_image_counter = 0  # 個別図形専用カウンター
        self.shape_metadata = shape_metadata  # 図形メタデータ出力フラグ
        self.output_format = output_format.lower() if output_format else 'png'
        
        # 出力形式の検証
        if self.output_format not in ('png', 'svg'):
            print(f"[WARNING] 不明な出力形式 '{output_format}'。'png'を使用します。")
            self.output_format = 'png'
        
        print(f"[INFO] 出力画像形式: {self.output_format.upper()}")
        
    def convert(self) -> str:
        """メイン変換処理"""
        print(f"[INFO] Word文書変換開始: {self.word_file}")
        
        # 1. 見出し構造を解析（参照リンク生成のため）
        self._analyze_headings()
        
        # 1.5. numbering定義を解析（デバッグ用）
        self._analyze_numbering_definitions()
        
        # 2. 文書本体を変換（Word文書の構造をそのまま再現）
        # 目次は文書内の適切な位置で挿入される
        self._convert_document_body()
        
        # 3. Markdownファイルを保存
        markdown_content = "\n".join(self.markdown_lines)
        output_file = os.path.join(self.output_dir, f"{self.base_name}.md")
        
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        
        print(f"[SUCCESS] 変換完了: {output_file}")
        return output_file
    
    def _analyze_headings(self):
        """見出し構造を解析"""
        print("[INFO] 見出し構造を解析中...")
        
        for paragraph in self.doc.paragraphs:
            style_name = paragraph.style.name.lower()
            text = paragraph.text.strip()
            
            if not text:
                continue
                
            # 見出しスタイルを検出
            heading_level = None
            if 'heading' in style_name:
                # "Heading 1" -> レベル1
                match = re.search(r'heading\s*(\d+)', style_name)
                if match:
                    heading_level = int(match.group(1))
            elif paragraph.style.base_style and 'heading' in paragraph.style.base_style.name.lower():
                # 基底スタイルから判定
                match = re.search(r'heading\s*(\d+)', paragraph.style.base_style.name.lower())
                if match:
                    heading_level = int(match.group(1))
            elif style_name in ['title']:
                heading_level = 1
            
            if heading_level:
                # アンカーIDを生成（日本語対応）
                anchor_id = self._generate_anchor_id(text)
                self.headings.append({
                    'level': heading_level,
                    'text': text,
                    'anchor': anchor_id
                })
                debug_print(f"[DEBUG] 見出し発見: レベル{heading_level} - {text}")
                
                # 章番号マッピングを構築（見出しテキストと段落の番号付け情報を使用）
                self._build_chapter_mapping(text, anchor_id, paragraph)
    
    def _generate_anchor_id(self, text: str) -> str:
        """アンカーIDを生成"""
        # 特殊文字を除去し、スペースをハイフンに
        anchor = re.sub(r'[^\w\s\u3040-\u309F\u30A0-\u30FF\u4E00-\u9FAF]', '', text)
        anchor = re.sub(r'\s+', '-', anchor)
        return anchor
    
    def _build_chapter_mapping(self, heading_text: str, anchor_id: str, paragraph=None):
        """見出しテキストから章番号マッピングを構築（Wordの自動番号付けに対応）"""
        import re
        
        # 段落の番号付け情報から章番号を取得
        auto_number = self._get_paragraph_auto_number(paragraph) if paragraph else None
        
        # スタイルベースの章番号も検出
        style_number = self._get_style_based_chapter_number(paragraph) if paragraph else None
        
        # 自動番号付けから章番号マッピングを作成
        chapter_number = auto_number or style_number
        if chapter_number:
            # 自動番号が取得できた場合のマッピング
            chapter_patterns = [
                f"{chapter_number}章",
                f"第{chapter_number}章",
                chapter_number
            ]
            
            # 階層的な章番号の場合、節やサブセクションのマッピングも追加
            if '.' in chapter_number:
                chapter_patterns.extend([
                    f"{chapter_number}節",
                    f"({chapter_number})",  # カッコ付きパターン
                ])
            
            for pattern in chapter_patterns:
                self.headings_map[pattern] = anchor_id
                # 見出しタイトルマッピングも作成
                if self.use_heading_text:
                    # 章番号部分を除去した見出しタイトルを取得
                    title_text = self._extract_heading_title(heading_text)
                    self.heading_titles_map[pattern] = title_text
                debug_print(f"[DEBUG] 自動章番号マッピング: '{pattern}' -> '#{anchor_id}'")
        
        # 見出しテキストから章番号を抽出
        patterns = [
            # 「第X章」パターン
            (r'第(\d+)章', lambda m: f"第{m.group(1)}章"),
            # 「X. 〜」パターン（「第X章」「X章」両方のマッピングを作成）
            (r'^(\d+)\.\s+', lambda m: (f"{m.group(1)}章", f"第{m.group(1)}章")),
            # 「X.Y 〜」パターン（「X.Y節」「X.Y」両方のマッピングを作成）
            (r'^(\d+\.\d+)\s+', lambda m: (m.group(1), f"{m.group(1)}節")),
            # 「X.Y.Z 〜」パターン  
            (r'^(\d+\.\d+\.\d+(?:\.\d+)*)\s+', lambda m: m.group(1)),
        ]
        
        for pattern, extractor in patterns:
            match = re.search(pattern, heading_text)
            if match:
                result = extractor(match)
                if isinstance(result, tuple):
                    # 複数のマッピングを作成
                    for chapter_ref in result:
                        self.headings_map[chapter_ref] = anchor_id
                        # 見出しタイトルマッピングも作成
                        if self.use_heading_text:
                            title_text = self._extract_heading_title(heading_text)
                            self.heading_titles_map[chapter_ref] = title_text
                        debug_print(f"[DEBUG] テキスト章番号マッピング: '{chapter_ref}' -> '#{anchor_id}'")
                else:
                    # 単一のマッピング
                    self.headings_map[result] = anchor_id
                    # 見出しタイトルマッピングも作成
                    if self.use_heading_text:
                        title_text = self._extract_heading_title(heading_text)
                        self.heading_titles_map[result] = title_text
                    debug_print(f"[DEBUG] テキスト章番号マッピング: '{result}' -> '#{anchor_id}'")
                break
    
    def _extract_heading_title(self, heading_text: str) -> str:
        """見出しテキストから章番号部分を除去してタイトル部分を抽出"""
        import re
        
        # 各種章番号パターンを除去
        patterns = [
            r'第\d+章\s*',        # 第X章
            r'\d+\.\s+',          # X.
            r'\d+\.\d+\s+',       # X.Y
            r'\d+\.\d+\.\d+\s+',  # X.Y.Z
            r'\d+\.\d+\.\d+\.\d+\s+',  # X.Y.Z.W
        ]
        
        result = heading_text
        for pattern in patterns:
            result = re.sub(pattern, '', result, count=1)  # 最初の一致のみ置換
        
        # 前後の空白を除去
        return result.strip()
    
    def _get_paragraph_auto_number(self, paragraph):
        """段落のWordの自動番号付けから章番号を取得"""
        try:
            # 段落の番号付け情報を取得
            numPr = paragraph._element.xpath('.//w:numPr')
            
            if not numPr:
                return None
                
            # numIdとilvlを取得
            numId_elem = paragraph._element.xpath('.//w:numPr/w:numId')
            ilvl_elem = paragraph._element.xpath('.//w:numPr/w:ilvl')
            
            if not numId_elem:
                return None
                
            numId = numId_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
            ilvl = int(ilvl_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if ilvl_elem else 0
            
            # numbering_types辞書を使用して番号付けスタイルを取得
            if hasattr(self, 'numbering_types') and numId in self.numbering_types:
                numbering_info = self.numbering_types[numId]
                if ilvl < len(numbering_info):
                    level_info = numbering_info[ilvl]
                    # 数値形式の場合のみ章番号として使用
                    if level_info.get('numFmt') == 'decimal':
                        # 実際の番号を計算（簡易実装）
                        return str(ilvl + 1)  # レベルベースの簡易計算
                        
            return None
            
        except Exception:
            return None
    
    def _get_style_based_chapter_number(self, paragraph):
        """スタイルベースの章番号を検出"""
        try:
            # 見出しレベルから章番号を推測
            style_name = paragraph.style.name.lower()
            
            # 見出しスタイルの場合、レベルから章番号を推測
            if 'heading' in style_name:
                import re
                match = re.search(r'heading\s*(\d+)', style_name)
                if match:
                    level = int(match.group(1))
                    if level <= 4:  # レベル1-4の見出しまで章番号として扱う
                        # 階層カウンターを初期化
                        if not hasattr(self, '_chapter_counters'):
                            self._chapter_counters = {}
                        
                        # 現在のレベル以下のカウンターをリセット
                        for l in range(level + 1, 5):
                            if l in self._chapter_counters:
                                self._chapter_counters[l] = 0
                        
                        # 現在のレベルのカウンターを増加
                        self._chapter_counters[level] = self._chapter_counters.get(level, 0) + 1
                        
                        # 階層的な章番号を生成
                        if level == 1:
                            return str(self._chapter_counters[1])
                        elif level == 2:
                            return f"{self._chapter_counters.get(1, 1)}.{self._chapter_counters[2]}"
                        elif level == 3:
                            return f"{self._chapter_counters.get(1, 1)}.{self._chapter_counters.get(2, 1)}.{self._chapter_counters[3]}"
                        elif level == 4:
                            return f"{self._chapter_counters.get(1, 1)}.{self._chapter_counters.get(2, 1)}.{self._chapter_counters.get(3, 1)}.{self._chapter_counters[4]}"
                        
            return None
            
        except Exception:
            return None
    
    def _has_toc_in_document(self):
        """Word文書内に目次フィールドが存在するかチェック"""
        try:
            # TOCフィールドやHyperlink要素をチェック
            for paragraph in self.doc.paragraphs:
                # 目次に関連するスタイルやテキストをチェック
                if paragraph.style.name.lower() in ['toc 1', 'toc 2', 'toc 3', 'toc heading', 'table of contents']:
                    return True
                # 「目次」「Contents」「Table of Contents」などのテキストをチェック
                text = paragraph.text.strip().lower()
                if text in ['目次', 'contents', 'table of contents', '目 次']:
                    return True
                # フィールドコードをチェック
                if 'TOC' in paragraph.text or 'HYPERLINK' in paragraph.text:
                    return True
            return False
        except Exception as e:
            debug_print(f"[DEBUG] 目次チェックエラー: {e}")
            return False
    
    def _generate_toc(self):
        """目次を生成"""
        print("[INFO] 目次を生成中...")
        
        self.markdown_lines.append("# 目次")
        self.markdown_lines.append("")
        
        for heading in self.headings:
            indent = "  " * (heading['level'] - 1)
            link = f"[{heading['text']}](#{heading['anchor']})"
            self.markdown_lines.append(f"{indent}- {link}")
        
        self.markdown_lines.append("")
        self.markdown_lines.append("---")
        self.markdown_lines.append("")
    
    def _convert_document_body(self):
        """文書本体を変換"""
        print("[INFO] 文書本体を変換中...")
        
        # 要素を順番に処理
        # 直前の要素を追跡するための変数
        previous_element_type = None
        
        for element in self.doc.element.body:
            if element.tag.endswith('}p'):  # 段落
                paragraph = self._find_paragraph_by_element(element)
                if paragraph:
                    self._convert_paragraph(paragraph)
                    # 段落に画像が含まれている場合は、その場で画像を処理
                    self._process_paragraph_images(paragraph)
                    # 見出しかどうかを記録
                    if self._is_heading(paragraph):
                        previous_element_type = 'heading'
                    elif self._is_list_item(paragraph):
                        previous_element_type = 'list'
                    else:
                        previous_element_type = 'paragraph'
            elif element.tag.endswith('}tbl'):  # 表
                table = self._find_table_by_element(element)
                if table:
                    # 見出しやリスト項目の直後にテーブルが来る場合は空行を挿入
                    if previous_element_type in ['heading', 'list']:
                        self.markdown_lines.append("")
                    self._convert_table(table)
                    previous_element_type = 'table'
        
        # 残りの画像を処理
        self._process_images()
    
    def _find_paragraph_by_element(self, element) -> Optional[Any]:
        """要素から段落オブジェクトを取得"""
        for paragraph in self.doc.paragraphs:
            if paragraph._element == element:
                return paragraph
        return None
    
    def _find_table_by_element(self, element) -> Optional[Any]:
        """要素から表オブジェクトを取得"""
        for table in self.doc.tables:
            if table._element == element:
                return table
        return None
    
    def _get_paragraph_text_without_hidden(self, paragraph) -> str:
        """段落から隠しテキスト（vanish属性を持つrun）を除外してテキストを取得"""
        text_parts = []
        for run in paragraph.runs:
            try:
                vanish_elem = run._element.xpath('.//w:vanish')
                if not vanish_elem:
                    if run.text:
                        text_parts.append(run.text)
            except Exception:
                if run.text:
                    text_parts.append(run.text)
        
        return ''.join(text_parts)
    
    def _convert_paragraph(self, paragraph):
        """段落を変換"""
        text = self._get_paragraph_text_without_hidden(paragraph).strip()
        style_name = paragraph.style.name.lower()
        
        if not text:
            # 空の段落は空行として処理
            self.markdown_lines.append("")
            return
        
        # Word文書内の目次を検出して展開
        if self._is_toc_placeholder(paragraph):
            if self.headings:
                print(f"[INFO] Word文書内の目次を検出、見出し{len(self.headings)}個で目次を展開します")
                self._generate_toc()
            else:
                print("[INFO] 見出しが見つからないため、目次プレースホルダーをスキップします")
            return
        
        # 見出しの処理
        if self._is_heading(paragraph):
            level = self._get_heading_level(paragraph)
            
            # 見出しレベルに応じてMarkdown記法を適用（アンカーIDなし）
            heading_prefix = "#" * level
            self.markdown_lines.append(f"{heading_prefix} {text}")
            self.markdown_lines.append("")
            return
        
        # リストの処理
        if self._is_list_item(paragraph):
            self._debug_paragraph_numbering(paragraph)  # デバッグ情報出力
            list_text = self._convert_list_item(paragraph)
            self.markdown_lines.append(list_text)
            return
        
        # 通常の段落（章番号参照をリンクに変換）
        text = self._convert_chapter_references(text)
        self.markdown_lines.append(text)
        self.markdown_lines.append("")
    
    def _is_toc_placeholder(self, paragraph) -> bool:
        """Word文書内の目次プレースホルダーかどうか判定"""
        try:
            # スタイル名での判定
            style_name = paragraph.style.name.lower()
            if style_name in ['toc 1', 'toc 2', 'toc 3', 'toc heading', 'table of contents']:
                debug_print(f"[DEBUG] 目次スタイル検出: {style_name}")
                return True
            
            # テキスト内容での判定
            text = paragraph.text.strip().lower()
            if text in ['目次', 'contents', 'table of contents', '目 次']:
                debug_print(f"[DEBUG] 目次テキスト検出: {text}")
                return True
            
            # フィールドコードでの判定
            if 'TOC' in paragraph.text or 'HYPERLINK' in paragraph.text:
                debug_print(f"[DEBUG] 目次フィールド検出: TOC/HYPERLINK")
                return True
            
            # Word文書のXML構造での判定
            for run in paragraph.runs:
                if run._element.xpath('.//w:fldChar[@w:fldCharType="begin"]'):
                    # フィールドの開始を検出
                    next_run = run._element.getnext()
                    while next_run is not None:
                        if 'TOC' in next_run.text if hasattr(next_run, 'text') else '':
                            debug_print(f"[DEBUG] Word目次フィールド検出")
                            return True
                        if next_run.xpath('.//w:fldChar[@w:fldCharType="end"]'):
                            break
                        next_run = next_run.getnext()
            
            return False
        except Exception as e:
            debug_print(f"[DEBUG] 目次判定エラー: {e}")
            return False
    
    def _is_heading(self, paragraph) -> bool:
        """見出しかどうか判定"""
        style_name = paragraph.style.name.lower()
        return ('heading' in style_name or 
                style_name in ['title'] or
                (paragraph.style.base_style and 'heading' in paragraph.style.base_style.name.lower()))
    
    def _get_heading_level(self, paragraph) -> int:
        """見出しレベルを取得"""
        style_name = paragraph.style.name.lower()
        
        if style_name == 'title':
            return 1
        
        match = re.search(r'heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        
        if paragraph.style.base_style:
            match = re.search(r'heading\s*(\d+)', paragraph.style.base_style.name.lower())
            if match:
                return int(match.group(1))
        
        return 1
    
    def _is_list_item(self, paragraph) -> bool:
        """リスト項目かどうか判定（段落番号と箇条書きを区別）"""
        # Word文書内のnumberingプロパティをチェック
        numPr = paragraph._element.xpath('.//w:numPr')
        if numPr:
            return True
        
        # スタイル名での判定
        style_name = paragraph.style.name.lower()
        if style_name in ['list paragraph', 'listparagraph']:
            return True
            
        # テキストの先頭パターンでの判定（フォールバック）
        text = paragraph.text.strip()
        return text.startswith(('•', '-', '*')) or re.match(r'^\d+\.', text)
    
    def _convert_chapter_references(self, text: str) -> str:
        """章番号参照をMarkdownリンクに変換（実際の見出しマッピングを使用）"""
        import re
        
        # すべてのパターンを一回のパスで処理
        def replace_chapter_ref(match):
            full_match = match.group(0)
            start_pos = match.start()
            end_pos = match.end()
            
            # 前後の文字を確認（文脈チェック）
            before = text[max(0, start_pos-1):start_pos] if start_pos > 0 else ''
            after = text[end_pos:end_pos+1] if end_pos < len(text) else ''
            
            # 「第X章」形式
            m = re.match(r'第(\d+)章', full_match)
            if m:
                chapter_ref = f"第{m.group(1)}章"
                if chapter_ref in self.headings_map:
                    # オプションに応じてリンクテキストを決定
                    link_text = self.heading_titles_map.get(chapter_ref, chapter_ref) if self.use_heading_text else chapter_ref
                    return f"[{link_text}](#{self.headings_map[chapter_ref]})"
                return full_match  # マッピングが存在しない場合は変換しない
            
            # 「X.Y.Z.W章」形式（複数階層）
            m = re.match(r'(\d+\.\d+(?:\.\d+)+)章', full_match)
            if m:
                chapter_ref = f"{m.group(1)}章"
                if chapter_ref in self.headings_map:
                    link_text = self.heading_titles_map.get(chapter_ref, chapter_ref) if self.use_heading_text else chapter_ref
                    return f"[{link_text}](#{self.headings_map[chapter_ref]})"
                return full_match
            
            # 「X.Y節」「X.Y.Z節」形式
            m = re.match(r'(\d+\.\d+(?:\.\d+)*)節', full_match)
            if m:
                section_ref = f"{m.group(1)}節"
                if section_ref in self.headings_map:
                    link_text = self.heading_titles_map.get(section_ref, section_ref) if self.use_heading_text else section_ref
                    return f"[{link_text}](#{self.headings_map[section_ref]})"
                return full_match
            
            # 「X章」形式
            m = re.match(r'(\d+)章', full_match)
            if m:
                chapter_ref = f"{m.group(1)}章"
                if chapter_ref in self.headings_map:
                    link_text = self.heading_titles_map.get(chapter_ref, chapter_ref) if self.use_heading_text else chapter_ref
                    return f"[{link_text}](#{self.headings_map[chapter_ref]})"
                return full_match
            
            # 「.X.Y.Z」形式
            m = re.match(r'\.(\d+\.\d+(?:\.\d+)*)', full_match)
            if m:
                section_ref = m.group(1)
                if section_ref in self.headings_map:
                    link_text = self.heading_titles_map.get(section_ref, section_ref) if self.use_heading_text else section_ref
                    return f"[.{link_text}](#{self.headings_map[section_ref]})"
                return full_match
            
            # 「(X.Y.Z)」形式（カッコ付き）
            m = re.match(r'\((\d+\.\d+(?:\.\d+)*)\)', full_match)
            if m:
                section_ref = m.group(1)
                if section_ref in self.headings_map:
                    link_text = self.heading_titles_map.get(section_ref, section_ref) if self.use_heading_text else section_ref
                    return f"[({link_text})](#{self.headings_map[section_ref]})"
                return full_match
            
            # 「X.Y.Z」形式（文脈チェックあり）
            m = re.match(r'(\d+\.\d+(?:\.\d+)*)', full_match)
            if m:
                # 適切な文脈かチェック（数字に囲まれていない、単語の境界にある）
                if (before in ['', ' ', '\n', '\t'] or before in ['、', '。', '：', '（', '(']) and \
                   (after in ['', ' ', '\n', '\t'] or after in ['、', '。', 'に', 'で', 'を', 'が', 'は', 'の', 'と', '）', ')']):
                    section_ref = m.group(1)
                    if section_ref in self.headings_map:
                        link_text = self.heading_titles_map.get(section_ref, section_ref) if self.use_heading_text else section_ref
                        return f"[{link_text}](#{self.headings_map[section_ref]})"
            
            return full_match
        
        # 章番号パターン
        patterns = [
            r'第\d+章',                    # 第X章
            r'\d+(?:\.\d+)+章',            # X.Y.Z章
            r'(?<!第)\d+章',               # X章（第がない）
            r'\d+(?:\.\d+)+節',            # X.Y節、X.Y.Z節
            r'\.\d+(?:\.\d+)+',            # .X.Y.Z
            r'\(\d+\.\d+(?:\.\d+)*\)',     # (X.Y.Z) カッコ付き
            r'\d+\.\d+(?:\.\d+)*'          # X.Y.Z（単独、後で文脈チェック）
        ]
        
        combined_pattern = '|'.join(f'({p})' for p in patterns)
        
        return re.sub(combined_pattern, replace_chapter_ref, text)

    def _convert_list_item(self, paragraph) -> str:
        """リスト項目を変換（段落番号と箇条書きを正しく区別）"""
        self._debug_paragraph_numbering(paragraph)
        
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Word文書のnumberingプロパティから判定
        numPr = paragraph._element.xpath('.//w:numPr')
        if numPr:
            # numIdとilvlを取得してリストタイプを判定
            numId_elem = paragraph._element.xpath('.//w:numPr/w:numId')
            ilvl_elem = paragraph._element.xpath('.//w:numPr/w:ilvl')
            
            numId = numId_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if numId_elem else '0'
            ilvl = int(ilvl_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) if ilvl_elem else 0
            
            # numbering_types辞書を使用して正確に判定
            if hasattr(self, 'numbering_types') and numId in self.numbering_types:
                numbering_info = self.numbering_types[numId]
                is_bullet = numbering_info['type'] == 'bullet'
                
                debug_print(f"[DEBUG] numId={numId} -> type={numbering_info['type']}, format='{numbering_info['format']}'")
            else:
                # フォールバック：従来の判定方法
                is_bullet = self._is_bullet_numbering(numId)
                debug_print(f"[DEBUG] numId={numId} -> フォールバック判定: {'bullet' if is_bullet else 'number'}")
            
            if is_bullet:
                # 箇条書きリスト
                indent = "  " * ilvl  # インデントレベル対応
                text = re.sub(r'^[•\-\*\d+\.]\s*', '', text)
                return f"{indent}- {text}"
            else:
                # 段落番号（番号付きリスト）
                indent = "  " * ilvl  # インデントレベル対応
                text = re.sub(r'^\d+\.\s*', '', text)
                return f"{indent}1. {text}"
        
        # フォールバック：テキストパターンで判定
        if re.match(r'^\d+\.', text):
            # 番号付きリスト
            text = re.sub(r'^\d+\.\s*', '', text)
            return f"1. {text}"
        else:
            # 箇条書きリスト
            text = re.sub(r'^[•\-\*]\s*', '', text)
            return f"- {text}"
    
    def _is_bullet_numbering(self, num_id: str) -> bool:
        """numbering IDから箇条書きかどうかを判定（改良版）"""
        try:
            # 一般的なパターンで判定
            # numId が奇数の場合は箇条書き、偶数の場合は番号付きという仮説
            # または特定のnumIdパターンで判定
            
            num_id_int = int(num_id)
            
            # 経験的パターン：
            # - numId 1,2,3 = 通常の番号付きリスト
            # - numId 4,5 = 機能説明の番号付きリスト  
            # - numId 13,20 = セクション見出し的な番号付きリスト
            # - 箇条書きは別のnumIdを使用する可能性
            
            # より正確な判定のため、実際のWord文書構造を詳しく見る必要がある
            # ここでは簡易判定：numId=1は箇条書きとして扱う
            if num_id_int == 1:
                return True
                
            return False
        except:
            return False
    
    def _debug_paragraph_numbering(self, paragraph):
        """段落の番号付け情報をデバッグ出力"""
        try:
            num_pr = paragraph._element.xpath('.//w:numPr')
            if num_pr:
                num_id_elem = paragraph._element.xpath('.//w:numPr/w:numId')
                ilvl_elem = paragraph._element.xpath('.//w:numPr/w:ilvl')
                
                num_id = num_id_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if num_id_elem else 'None'
                ilvl = ilvl_elem[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if ilvl_elem else 'None'
                
                text = paragraph.text.strip()[:30]
                debug_print(f"[DEBUG] リスト項目: '{text}' | numId={num_id} | ilvl={ilvl}")
        except Exception as e:
            debug_print(f"[DEBUG] 番号付けデバッグエラー: {e}")
    
    def _analyze_numbering_definitions(self):
        """numbering.xmlから番号付け定義を解析"""
        try:
            # numbering_types辞書を初期化
            self.numbering_types = {}
            
            # Word文書からnumbering.xmlを取得
            numbering_part = None
            for rel in self.doc.part.rels.values():
                if 'numbering' in rel.reltype:
                    numbering_part = rel.target_part
                    break
            
            if numbering_part:
                numbering_xml = numbering_part.blob.decode('utf-8')
                debug_print(f"[DEBUG] numbering.xml の一部: {numbering_xml[:500]}")
                
                # 各numIdのlvlText（表示形式）を解析
                import xml.etree.ElementTree as ET
                root = ET.fromstring(numbering_xml)
                
                # numId -> abstractNumId のマッピングを作成
                num_to_abstract = {}
                for num in root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num'):
                    num_id = num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId')
                    abstract_num_id = num.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                    if abstract_num_id is not None:
                        abstract_id = abstract_num_id.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                        num_to_abstract[num_id] = abstract_id
                        debug_print(f"[DEBUG] numId={num_id} -> abstractNumId={abstract_id}")
                
                # abstractNum定義から実際の番号形式を解析
                for abstract_num in root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum'):
                    abstract_id = abstract_num.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId')
                    
                    # レベル0の番号形式を取得（名前空間を考慮）
                    lvl_elements = abstract_num.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lvl')
                    lvl_element = None
                    
                    # ilvl="0"のlvl要素を探す
                    for lvl in lvl_elements:
                        ilvl_val = lvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                        if ilvl_val == '0':
                            lvl_element = lvl
                            break
                    
                    if lvl_element is not None:
                        lvl_text = lvl_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lvlText')
                        num_fmt = lvl_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt')
                        
                        format_text = lvl_text.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if lvl_text is not None else ''
                        format_type = num_fmt.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if num_fmt is not None else ''
                        
                        # 番号形式を判定
                        is_bullet = False
                        if format_type == 'bullet' or format_text in ['·', '•', '-', '○', '■', 'l']:
                            is_bullet = True
                        elif '%1' in format_text and format_type in ['decimal', 'lowerLetter', 'upperLetter']:
                            is_bullet = False
                        
                        # 該当するnumIdに情報を保存
                        for num_id, mapped_abstract_id in num_to_abstract.items():
                            if mapped_abstract_id == abstract_id:
                                self.numbering_types[num_id] = {
                                    'type': 'bullet' if is_bullet else 'number',
                                    'format': format_text,
                                    'format_type': format_type,
                                    'abstract_id': abstract_id
                                }
                                debug_print(f"[DEBUG] numId={num_id}: type={'bullet' if is_bullet else 'number'}, format='{format_text}', format_type='{format_type}'")
                        
        except Exception as e:
            debug_print(f"[DEBUG] numbering解析エラー: {e}")
            import traceback
            traceback.print_exc()
    
    def _convert_table(self, table):
        """表を変換"""
        print("[INFO] 表を変換中...")
        
        if not table.rows:
            self.markdown_lines.append("*空の表*")
            self.markdown_lines.append("")
            return
        
        # ヘッダー行
        header_row = table.rows[0]
        header_cells = [self._process_table_cell_text(cell) for cell in header_row.cells]
        
        # Markdownテーブル形式で出力
        self.markdown_lines.append("| " + " | ".join(header_cells) + " |")
        self.markdown_lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
        
        # データ行
        for row in table.rows[1:]:
            cells = [self._process_table_cell_text(cell) for cell in row.cells]
            # セル数を調整
            while len(cells) < len(header_cells):
                cells.append("")
            cells = cells[:len(header_cells)]
            
            self.markdown_lines.append("| " + " | ".join(cells) + " |")
        
        # 表の後に空行を追加（次の要素との間隔確保）
        self.markdown_lines.append("")
        
        self.markdown_lines.append("")
    
    def _process_table_cell_text(self, cell):
        """表のセル内のテキストを処理（改行を<br>に変換）"""
        # セル内の段落を取得
        paragraphs = cell.paragraphs
        
        if not paragraphs:
            return ""
        
        # 段落ごとのテキストを取得
        paragraph_texts = []
        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if text:  # 空でない段落のみを処理
                paragraph_texts.append(text)
        
        # 段落間を<br>で結合
        cell_text = "<br>".join(paragraph_texts)
        
        # 章番号のリンク変換を適用
        cell_text = self._convert_chapter_references(cell_text)
        
        # Markdownテーブル内で問題となる文字をエスケープ
        # ただし、リンク内のパイプは除外
        escaped_text = ""
        in_link = False
        i = 0
        while i < len(cell_text):
            if cell_text[i:i+1] == '[':
                in_link = True
                escaped_text += cell_text[i]
            elif cell_text[i:i+1] == ')' and in_link:
                in_link = False
                escaped_text += cell_text[i]
            elif cell_text[i:i+1] == '|' and not in_link:
                escaped_text += '\\|'
            else:
                escaped_text += cell_text[i]
            i += 1
        
        return escaped_text
    
    def _process_paragraph_images(self, paragraph):
        """段落内の画像を処理"""
        
        # Word図形キャンバスがある場合は複合図形として処理
        # 処理が行われた場合のみ早期終了、そうでなければ通常の画像処理にフォールバック
        if self._has_word_processing_canvas(paragraph):
            if self._process_composite_figure(paragraph):
                return
        
        # 段落内のRunを調べて画像があるかチェック
        for run in paragraph.runs:
            # drawing要素を取得
            drawings = run._element.xpath('.//w:drawing')
            for drawing in drawings:
                for inline_shape in drawing.xpath('.//a:blip'):
                    # 画像の参照IDを取得
                    embed_id = inline_shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id:
                        # 対応するリレーションを探す
                        for rel in paragraph.part.rels.values():
                            if rel.rId == embed_id and "image" in rel.reltype:
                                # その場で画像を処理（drawing要素も渡す）
                                self._extract_and_convert_image_inline(rel, drawing)
                                return  # 1つの段落につき1つの画像のみ処理
    
    def _extract_and_convert_image_inline(self, rel, drawing_element=None):
        """インライン画像を抽出・変換"""
        try:
            # 参照されている画像として記録
            self.referenced_images.add(rel.rId)
            
            # すでに処理済みの画像はスキップ（ハッシュベース重複検出）
            image_data = rel.target_part.blob
            image_hash = hash(image_data)
            
            # 重複チェック（rIdまたはハッシュでチェック）
            if rel.rId in self.processed_images or any(info['hash'] == image_hash for info in self.processed_images.values()):
                return
            
            # 処理済みに記録
            self.processed_images[rel.rId] = {
                'hash': image_hash,
                'filename': None,  # 後で設定
                'path': None  # 後で設定
            }
            
            # 画像形式を判定
            extension = self._detect_image_format(image_data, rel.target_ref)
            
            # ファイル名を生成（専用カウンター使用）
            self.regular_image_counter += 1
            image_filename = f"{self.base_name}_image_{self.regular_image_counter:03d}{extension}"
            image_path = os.path.join(self.images_dir, image_filename)
            
            # EMF/WMFの場合は変換
            if extension in ['.emf', '.wmf']:
                converted_path = self._convert_vector_image(image_data, image_path)
                if converted_path:
                    image_path = converted_path
                    image_filename = os.path.basename(converted_path)
            else:
                # そのまま保存
                with open(image_path, 'wb') as f:
                    f.write(image_data)
            
            # 処理済み情報を更新
            self.processed_images[rel.rId]['filename'] = image_filename
            self.processed_images[rel.rId]['path'] = image_path
            
            # Markdownに追加（ファイル名をURLエンコード）
            encoded_filename = urllib.parse.quote(image_filename)
            self.markdown_lines.append(f"![](images/{encoded_filename})")
            self.markdown_lines.append("")
            
            if self.shape_metadata and drawing_element is not None:
                try:
                    metadata = self._extract_shape_metadata_from_drawing(drawing_element)
                    if metadata.get('shapes'):
                        text_metadata = self._format_shape_metadata_as_text(metadata)
                        json_metadata = self._format_shape_metadata_as_json(metadata)
                        
                        if text_metadata:
                            self.markdown_lines.append("")
                            self.markdown_lines.append(text_metadata)
                            self.markdown_lines.append("")
                        
                        if json_metadata and json_metadata != "{}":
                            self.markdown_lines.append("<details>")
                            self.markdown_lines.append("<summary>JSON形式の図形情報</summary>")
                            self.markdown_lines.append("")
                            self.markdown_lines.append("```json")
                            self.markdown_lines.append(json_metadata)
                            self.markdown_lines.append("```")
                            self.markdown_lines.append("")
                            self.markdown_lines.append("</details>")
                            self.markdown_lines.append("")
                        
                        debug_print(f"[DEBUG] 図形メタデータ追加: {len(metadata['shapes'])} shapes")
                except Exception as e:
                    print(f"[WARNING] 図形メタデータ追加失敗: {e}")
            
            print(f"[SUCCESS] 画像をインライン処理: {image_filename}")
            
        except Exception as e:
            print(f"[ERROR] インライン画像処理エラー: {e}")
    
    
    def _process_images(self):
        """実際に文書内で参照されている画像のみを処理（文書末尾の画像など）- 重複除去"""
        print("[INFO] 残りの画像を処理中...")
        
        # 実際に参照されている画像のみを処理（重複チェック強化）
        for rel in self.doc.part.rels.values():
            if ("image" in rel.reltype and 
                rel.rId in self.referenced_images and 
                rel.rId not in self.processed_images):
                
                # 画像データのハッシュも確認して真の重複を防ぐ
                image_data = rel.target_part.blob
                image_hash = hash(image_data)
                
                # ハッシュベースでも重複チェック
                already_processed = False
                for processed_rel_id, processed_info in self.processed_images.items():
                    if isinstance(processed_info, dict) and processed_info.get('hash') == image_hash:
                        already_processed = True
                        debug_print(f"[DEBUG] 画像重複スキップ: {rel.rId} (ハッシュ重複)")
                        break
                
                if not already_processed:
                    self._extract_and_convert_image(rel)
    
    def _extract_and_convert_image(self, rel):
        """画像を抽出・変換（重複防止強化）"""
        try:
            image_data = rel.target_part.blob
            image_hash = hash(image_data)
            
            # 画像形式を判定
            extension = self._detect_image_format(image_data, rel.target_ref)
            
            # ファイル名を生成（専用カウンター使用）
            self.regular_image_counter += 1
            image_filename = f"{self.base_name}_image_{self.regular_image_counter:03d}{extension}"
            image_path = os.path.join(self.images_dir, image_filename)
            
            # EMF/WMFの場合は変換
            if extension in ['.emf', '.wmf']:
                converted_path = self._convert_vector_image(image_data, image_path)
                if converted_path:
                    image_path = converted_path
                    image_filename = os.path.basename(converted_path)
            else:
                # そのまま保存
                with open(image_path, 'wb') as f:
                    f.write(image_data)
            
            # Markdownに追加（ファイル名をURLエンコード）
            encoded_filename = urllib.parse.quote(image_filename)
            self.markdown_lines.append(f"![](images/{encoded_filename})")
            self.markdown_lines.append("")
            
            # 処理済みとして記録（ハッシュ付き）
            self.processed_images[rel.rId] = {
                'filename': image_filename,
                'hash': image_hash,
                'path': image_path
            }
            
            print(f"[SUCCESS] 画像を処理: {image_filename}")
            
        except Exception as e:
            print(f"[ERROR] 画像処理エラー: {e}")
    
    def _has_word_processing_canvas(self, paragraph):
        """段落にWord図形キャンバス（またはグループ、または個別図形）があるかチェック"""
        try:
            drawings = paragraph._element.xpath('.//w:drawing')
            for drawing in drawings:
                # Word Processing Canvas (wpc) をチェック
                canvas_elements = drawing.xpath('.//*[local-name()="wpc"]')
                if canvas_elements:
                    debug_print("[DEBUG] Word Processing Canvas検出")
                    return True
                
                # Word Processing Group (wpg) をチェック
                group_elements = drawing.xpath('.//*[local-name()="wgp"]')
                if group_elements:
                    debug_print("[DEBUG] Word Processing Group検出")
                    return True
                
                # 個別のWord図形 (wps:wsp) をチェック（グループに含まれないもの）
                shape_elements = drawing.xpath('.//*[local-name()="wsp"]')
                if shape_elements:
                    debug_print("[DEBUG] Word Processing Shape検出")
                    return True
                    
            return False
        except Exception as e:
            print(f"[ERROR] キャンバス検出エラー: {e}")
            return False
    
    def _process_composite_figure(self, paragraph):
        """Word図形キャンバス/グループ/個別図形から複合図形を処理
        
        段落単位で図形を分類し、以下のルールで処理:
        1. wpg/wpcがある段落では、グループのみを処理（個別wspは無視）
        2. wspのみの段落では、すべてのdrawingを1つの画像にまとめる
        3. pic（通常の画像）がある段落では、段落グループ化をスキップ
        
        Returns:
            bool: 処理が行われた場合はTrue、スキップした場合はFalse
        """
        try:
            print("[INFO] 複合図形を処理中...")
            
            # Drawing要素を取得
            drawings = paragraph._element.xpath('.//w:drawing')
            if not drawings:
                return False
            
            # Drawing要素を分類
            canvas_drawings = []  # wpc/wpgを含むdrawing
            shape_only_drawings = []  # wspのみを含むdrawing
            has_picture = False  # pic（通常の画像）があるかどうか
            
            for drawing in drawings:
                # 1. Word Processing Canvas (wpc) をチェック（最優先）
                # wpg/wpcは内部にpicを含む場合があるため、picより先にチェック
                canvas_elements = drawing.xpath('.//*[local-name()="wpc"]')
                if canvas_elements:
                    canvas_drawings.append((drawing, canvas_elements[0], 'wpc'))
                    continue
                
                # 2. Word Processing Group (wpg) をチェック
                group_elements = drawing.xpath('.//*[local-name()="wgp"]')
                if group_elements:
                    canvas_drawings.append((drawing, group_elements[0], 'wpg'))
                    continue
                
                # 3. 個別のWord図形 (wps:wsp) をチェック
                shape_elements = drawing.xpath('.//*[local-name()="wsp"]')
                if shape_elements:
                    shape_only_drawings.append(drawing)
                    continue
                
                # 4. pic（通常の画像）をチェック - blip参照を持つ画像
                # wpg/wpc/wspに含まれないdrawingのみがここに到達
                pic_elements = drawing.xpath('.//*[local-name()="pic"]')
                blip_elements = drawing.xpath('.//*[local-name()="blip"]')
                if pic_elements or blip_elements:
                    has_picture = True
                    debug_print("[DEBUG] 段落内にpic（通常の画像）を検出")
            
            # wpc/wpgがある場合は、それらのみを処理（個別wspは無視）
            if canvas_drawings:
                processed = False
                for drawing, element, element_type in canvas_drawings:
                    print(f"[INFO] Word Processing {element_type.upper()} として処理")
                    if self._process_canvas_as_vector(element, drawing):
                        print("[SUCCESS] ベクター処理成功")
                        processed = True
                    else:
                        print("[ERROR] ベクター処理失敗")
                return processed
            
            # pic（通常の画像）がある段落では、段落グループ化をスキップ
            # 通常の画像処理ロジックに任せる
            if has_picture:
                debug_print("[INFO] 段落内にpic（通常の画像）があるため、段落グループ化をスキップ")
                return False
            
            # wspのみの段落では、すべてのdrawingを1つの画像にまとめる
            if shape_only_drawings:
                if len(shape_only_drawings) == 1:
                    # 1つだけの場合は従来通り処理
                    print("[INFO] 単一のWord Processing Shape として処理")
                    if self._process_shape_as_vector(None, shape_only_drawings[0]):
                        print("[SUCCESS] 個別図形ベクター処理成功")
                        return True
                    else:
                        print("[ERROR] 個別図形ベクター処理失敗")
                        return False
                else:
                    # 複数の場合は1つの画像にまとめる
                    print(f"[INFO] {len(shape_only_drawings)}個の個別図形を1つの画像にまとめて処理")
                    if self._process_shape_cluster_as_vector(shape_only_drawings):
                        print("[SUCCESS] 図形クラスターベクター処理成功")
                        return True
                    else:
                        print("[ERROR] 図形クラスターベクター処理失敗")
                        return False
            
            return False
            
        except Exception as e:
            print(f"[ERROR] 複合図形処理エラー: {e}")
            return False
    
    def _process_shape_as_vector(self, shape_element, drawing_element):
        """個別のWord図形をベクター画像として処理"""
        try:
            debug_print("[INFO] 個別図形をベクター画像として処理中...")
            
            # 一時的なWord文書を作成して図形のみを含める
            temp_doc_path = self._create_canvas_document(shape_element, drawing_element)
            if not temp_doc_path:
                return False
            
            debug_print(f"[DEBUG] 一時Word文書作成: {temp_doc_path}")
            
            # LibreOfficeでPDFに変換
            temp_pdf_path = self._convert_document_to_pdf(temp_doc_path)
            if not temp_doc_path:
                os.unlink(temp_doc_path)
                return False
            
            debug_print(f"[DEBUG] PDF変換完了: {temp_pdf_path}")
            
            # PDFから画像に変換（個別図形用カウンターを使用）
            self.shape_image_counter += 1
            ext = self.output_format
            image_filename = f"{self.base_name}_shape_{self.shape_image_counter:03d}.{ext}"
            image_path = os.path.join(self.images_dir, image_filename)
            
            # 出力形式に応じて変換
            if self.output_format == 'svg':
                convert_success = self._convert_pdf_to_svg(temp_pdf_path, image_path)
            else:
                convert_success = self._convert_pdf_to_png(temp_pdf_path, image_path)
            
            if convert_success:
                # Markdownに追加（ファイル名をURLエンコード）
                encoded_filename = urllib.parse.quote(image_filename)
                self.markdown_lines.append(f"![](images/{encoded_filename})")
                self.markdown_lines.append("")
                
                debug_print(f"[SUCCESS] 個別図形を処理: {image_filename}")
                
                # 一時ファイルを削除
                os.unlink(temp_doc_path)
                os.unlink(temp_pdf_path)
                return True
            
            # 一時ファイルを削除
            os.unlink(temp_doc_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            return False
            
        except Exception as e:
            print(f"[ERROR] 個別図形処理エラー: {e}")
            return False
    
    def _process_shape_cluster_as_vector(self, drawing_elements):
        """複数の個別図形を1つのベクター画像として処理
        
        同じ段落内の複数のdrawing要素を1つの画像にまとめる
        """
        try:
            debug_print(f"[INFO] {len(drawing_elements)}個の図形を1つの画像にまとめて処理中...")
            
            # 一時的なWord文書を作成して複数の図形を含める
            temp_doc_path = self._create_canvas_document(None, drawing_elements)
            if not temp_doc_path:
                return False
            
            debug_print(f"[DEBUG] 一時Word文書作成: {temp_doc_path}")
            
            # LibreOfficeでPDFに変換
            temp_pdf_path = self._convert_document_to_pdf(temp_doc_path)
            if not temp_pdf_path:
                os.unlink(temp_doc_path)
                return False
            
            debug_print(f"[DEBUG] PDF変換完了: {temp_pdf_path}")
            
            # PDFから画像に変換（図形クラスター用カウンターを使用）
            self.shape_image_counter += 1
            ext = self.output_format
            image_filename = f"{self.base_name}_shape_{self.shape_image_counter:03d}.{ext}"
            image_path = os.path.join(self.images_dir, image_filename)
            
            # 出力形式に応じて変換
            if self.output_format == 'svg':
                convert_success = self._convert_pdf_to_svg(temp_pdf_path, image_path)
            else:
                convert_success = self._convert_pdf_to_png(temp_pdf_path, image_path)
            
            if convert_success:
                # Markdownに追加（ファイル名をURLエンコード）
                encoded_filename = urllib.parse.quote(image_filename)
                self.markdown_lines.append(f"![](images/{encoded_filename})")
                self.markdown_lines.append("")
                
                debug_print(f"[SUCCESS] 図形クラスターを処理: {image_filename}")
                
                # 一時ファイルを削除
                os.unlink(temp_doc_path)
                os.unlink(temp_pdf_path)
                return True
            
            # 一時ファイルを削除
            os.unlink(temp_doc_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            return False
            
        except Exception as e:
            print(f"[ERROR] 図形クラスター処理エラー: {e}")
            return False
    
    def _process_canvas_as_vector(self, canvas_element, drawing_element):
        """Word図形キャンバス全体をベクター画像として処理"""
        try:
            print("[INFO] キャンバス全体をベクター画像として処理中...")
            
            # 一時的なWord文書を作成してキャンバスのみを含める
            temp_doc_path = self._create_canvas_document(canvas_element, drawing_element)
            if not temp_doc_path:
                return False
            
            debug_print(f"[DEBUG] 一時Word文書作成: {temp_doc_path}")
            
            # LibreOfficeでPDFに変換
            temp_pdf_path = self._convert_document_to_pdf(temp_doc_path)
            if not temp_pdf_path:
                os.unlink(temp_doc_path)
                return False
            
            debug_print(f"[DEBUG] PDF変換完了: {temp_pdf_path}")
            
            # PDFの内容を確認
            self._debug_pdf_content(temp_pdf_path)
            
            # PDFから画像に変換（専用カウンターを使用、出力形式に応じてPNGまたはSVG）
            self.vector_image_counter += 1
            ext = self.output_format
            image_filename = f"{self.base_name}_vector_composite_{self.vector_image_counter:03d}.{ext}"
            image_path = os.path.join(self.images_dir, image_filename)
            
            # 出力形式に応じて変換
            if self.output_format == 'svg':
                convert_success = self._convert_pdf_to_svg(temp_pdf_path, image_path)
            else:
                convert_success = self._convert_pdf_to_png(temp_pdf_path, image_path)
            
            if convert_success:
                # 生成された画像の詳細を確認
                self._debug_image_info(image_path)
                
                # Markdownに追加（ファイル名をURLエンコード）
                encoded_filename = urllib.parse.quote(image_filename)
                self.markdown_lines.append(f"![](images/{encoded_filename})")
                self.markdown_lines.append("")
                
                if self.shape_metadata:
                    try:
                        metadata = self._extract_shape_metadata_from_drawing(drawing_element)
                        if metadata.get('shapes'):
                            text_metadata = self._format_shape_metadata_as_text(metadata)
                            json_metadata = self._format_shape_metadata_as_json(metadata)
                            
                            if text_metadata:
                                self.markdown_lines.append("")
                                self.markdown_lines.append(text_metadata)
                                self.markdown_lines.append("")
                            
                            if json_metadata and json_metadata != "{}":
                                self.markdown_lines.append("<details>")
                                self.markdown_lines.append("<summary>JSON形式の図形情報</summary>")
                                self.markdown_lines.append("")
                                self.markdown_lines.append("```json")
                                self.markdown_lines.append(json_metadata)
                                self.markdown_lines.append("```")
                                self.markdown_lines.append("")
                                self.markdown_lines.append("</details>")
                                self.markdown_lines.append("")
                            
                            debug_print(f"[DEBUG] 図形メタデータ追加: {len(metadata['shapes'])} shapes")
                    except Exception as e:
                        print(f"[WARNING] 図形メタデータ追加失敗: {e}")
                
                print(f"[SUCCESS] ベクター複合図形を処理: {image_filename}")
                
                # デバッグ用にPDFも保存
                debug_pdf_path = os.path.join('output/debug', f"{os.path.splitext(os.path.basename(image_filename))[0]}.pdf")
                os.makedirs('output/debug', exist_ok=True)
                shutil.copy2(temp_pdf_path, debug_pdf_path)
                debug_print(f"[DEBUG] PDFデバッグファイル保存: {debug_pdf_path}")
                
                # 一時ファイルを削除
                os.unlink(temp_doc_path)
                os.unlink(temp_pdf_path)
                return True
            
            # 一時ファイルを削除
            os.unlink(temp_doc_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            return False
            
        except Exception as e:
            print(f"[ERROR] ベクター画像処理エラー: {e}")
            return False
    
    def _build_theme_color_map(self, theme_data):
        """テーマデータからカラーマップを構築
        
        Args:
            theme_data: テーマXMLのバイナリデータ
            
        Returns:
            dict: テーマ色名からRGB値へのマッピング（例: {'lt1': 'FFFFFF', 'dk1': '000000'}）
        """
        color_map = {}
        if not theme_data:
            return color_map
        
        try:
            from lxml import etree
            theme_tree = etree.fromstring(theme_data)
            
            # カラースキームを探す
            clr_scheme = theme_tree.xpath('.//*[local-name()="clrScheme"]')
            if not clr_scheme:
                return color_map
            
            for child in clr_scheme[0]:
                # タグ名からテーマ色名を取得（例: dk1, lt1, accent1）
                tag_name = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                
                # srgbClrを探す
                srgb_elems = child.xpath('.//*[local-name()="srgbClr"]/@val')
                if srgb_elems:
                    color_map[tag_name] = srgb_elems[0]
                    continue
                
                # sysClrを探す（lastClr属性にRGB値がある）
                sys_clr_elems = child.xpath('.//*[local-name()="sysClr"]')
                if sys_clr_elems:
                    last_clr = sys_clr_elems[0].get('lastClr')
                    if last_clr:
                        color_map[tag_name] = last_clr
            
            debug_print(f"[DEBUG] テーマカラーマップ構築: {len(color_map)}色")
        except Exception as e:
            debug_print(f"[DEBUG] テーマカラーマップ構築エラー: {e}")
        
        return color_map
    
    def _convert_text_scheme_colors(self, drawing_element, theme_color_map):
        """drawing要素内のテキスト色のschemeClrをsrgbClrに変換
        
        fontRef内のschemeClrを変換する（テキストのフォント色を定義）
        
        Args:
            drawing_element: drawing要素
            theme_color_map: テーマ色名からRGB値へのマッピング
            
        Returns:
            str: 変換後のXML文字列
        """
        if not theme_color_map:
            return ET.tostring(drawing_element, encoding='unicode')
        
        try:
            from lxml import etree
            xml_str = ET.tostring(drawing_element, encoding='unicode')
            tree = etree.fromstring(xml_str.encode('utf-8'))
            
            ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
            converted_count = 0
            
            # fontRef内のschemeClrを変換（テキスト色）
            font_ref_schemes = tree.xpath(
                './/*[local-name()="fontRef"]/*[local-name()="schemeClr"]'
            )
            
            for scheme_clr in font_ref_schemes:
                val = scheme_clr.get('val')
                if val and val in theme_color_map:
                    rgb_val = theme_color_map[val]
                    font_ref = scheme_clr.getparent()
                    
                    # srgbClr要素を作成
                    new_srgb = etree.Element(f"{{{ns_a}}}srgbClr", val=rgb_val)
                    
                    # 子要素をコピー
                    for child in list(scheme_clr):
                        new_srgb.append(child)
                    
                    # 置換
                    font_ref.remove(scheme_clr)
                    font_ref.append(new_srgb)
                    converted_count += 1
            
            if converted_count > 0:
                debug_print(f"[DEBUG] テキスト色変換: {converted_count}箇所")
            
            return etree.tostring(tree, encoding='unicode')
            
        except Exception as e:
            debug_print(f"[DEBUG] テキスト色変換エラー: {e}")
            return ET.tostring(drawing_element, encoding='unicode')
    
    def _create_canvas_document(self, canvas_element, drawing_elements):
        """キャンバス要素のみを含む一時Word文書を作成
        
        Args:
            canvas_element: キャンバス要素（互換性のため維持、現在は使用しない）
            drawing_elements: drawing要素（単一またはリスト）
        """
        try:
            debug_print("[DEBUG] Word文書作成開始...")
            debug_print("[DEBUG] SCHEMECLR_PATCH_V3_20251211")
            
            # drawing_elementsがリストでない場合は単一要素として扱う
            if not isinstance(drawing_elements, (list, tuple)):
                drawing_elements = [drawing_elements]
            
            # 元の文書からリレーション情報を取得
            original_rels = {}
            theme_data = None
            theme_rel_id = None
            try:
                for rel in self.doc.part.rels.values():
                    if "image" in rel.reltype:
                        original_rels[rel.rId] = rel.target_part.blob
                    # テーマリレーションを取得（schemeClr参照の解決に必要）
                    elif "theme" in rel.reltype:
                        try:
                            theme_data = rel.target_part.blob
                            theme_rel_id = rel.rId
                            debug_print(f"[DEBUG] テーマ取得: {rel.rId}")
                        except Exception as theme_error:
                            debug_print(f"[DEBUG] テーマ取得エラー: {theme_error}")
                debug_print(f"[DEBUG] 取得したリレーション数: {len(original_rels)}, テーマ: {theme_rel_id is not None}")
            except Exception as rel_error:
                debug_print(f"[DEBUG] リレーション取得エラー: {rel_error}")
            
            # テーマからカラーマップを作成（schemeClr→srgbClr変換用）
            theme_color_map = self._build_theme_color_map(theme_data)
            
            # 複数のdrawing XMLを連結し、テキスト色のschemeClrをsrgbClrに変換
            converted_drawings = []
            for d in drawing_elements:
                converted_xml = self._convert_text_scheme_colors(d, theme_color_map)
                converted_drawings.append(converted_xml)
            drawings_xml = "".join(converted_drawings)
            debug_print(f"[DEBUG] Drawing XML長: {len(drawings_xml)}")
            
            # より適切なWord文書XMLを作成
            doc_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
            xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
            xmlns:w10="urn:schemas-microsoft-com:office:word"
            xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
            xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
            xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
            xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
            xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
    <w:body>
        <w:p>
            <w:r>
                {drawings_xml}
            </w:r>
        </w:p>
        <w:sectPr>
            <w:pgSz w:w="11906" w:h="16838"/>
            <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>
        </w:sectPr>
    </w:body>
</w:document>'''

            # 一時ファイルに保存
            temp_docx_path = tempfile.mktemp(suffix='.docx')
            
            # Word文書ZIPファイルを作成
            with zipfile.ZipFile(temp_docx_path, 'w', zipfile.ZIP_DEFLATED) as docx:
                # Content_Types.xml（テーマがある場合はOverrideを追加）
                content_types_base = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Default Extension="png" ContentType="image/png"/>
    <Default Extension="jpeg" ContentType="image/jpeg"/>
    <Default Extension="jpg" ContentType="image/jpeg"/>
    <Default Extension="emf" ContentType="image/x-emf"/>
    <Default Extension="wmf" ContentType="image/x-wmf"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'''
                
                # テーマがある場合はOverrideを追加
                if theme_data:
                    content_types_base += '''
    <Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>'''
                
                content_types_base += '''
</Types>'''
                docx.writestr('[Content_Types].xml', content_types_base)
                
                # メインリレーション
                docx.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>''')
                
                # 文書リレーション（画像がある場合）
                doc_rels_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'''
                
                # 画像リレーションを追加
                for i, (rel_id, image_data) in enumerate(original_rels.items(), 1):
                    extension = self._detect_image_format(image_data, '')
                    target = f"media/image{i}{extension}"
                    doc_rels_xml += f'''
    <Relationship Id="{rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="{target}"/>'''
                    
                    # 画像ファイルを追加
                    docx.writestr(f"word/{target}", image_data)
                
                # テーマリレーションを追加（schemeClr参照の解決に必要）
                if theme_data and theme_rel_id:
                    doc_rels_xml += f'''
    <Relationship Id="{theme_rel_id}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/>'''
                    # テーマファイルを追加
                    docx.writestr('word/theme/theme1.xml', theme_data)
                    debug_print(f"[DEBUG] テーマファイル追加: word/theme/theme1.xml")
                
                doc_rels_xml += '''
</Relationships>'''
                
                docx.writestr('word/_rels/document.xml.rels', doc_rels_xml)
                
                # 文書内容
                docx.writestr('word/document.xml', doc_xml)
            
            print(f"[INFO] 一時Word文書作成完了: {temp_docx_path}")
            
            # デバッグ用に一時Word文書を保存
            debug_docx_dir = os.path.join('output', 'debug')
            os.makedirs(debug_docx_dir, exist_ok=True)
            debug_docx_path = os.path.join(debug_docx_dir, f"temp_docx_{os.path.basename(temp_docx_path)}")
            shutil.copy2(temp_docx_path, debug_docx_path)
            debug_print(f"[DEBUG] 一時Word文書デバッグコピー: {debug_docx_path}")
            
            return temp_docx_path
            
        except Exception as e:
            print(f"[ERROR] 一時文書作成エラー: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _convert_document_to_pdf(self, docx_path):
        """Word文書をPDFに変換（最高品質設定）"""
        try:
            temp_dir = tempfile.mkdtemp()
            
            # LibreOfficeでPDFに変換（最高品質設定）
            cmd = [
                LIBREOFFICE_PATH,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                docx_path
            ]
            
            # 環境変数でPDF品質を最高設定に
            env = os.environ.copy()
            env['SAL_DISABLE_OPENCL'] = '1'  # OpenCLを無効化して安定性を向上
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, env=env)
            
            if result.returncode == 0:
                # 変換されたPDFのパスを探す
                for file in os.listdir(temp_dir):
                    if file.endswith('.pdf'):
                        pdf_path = os.path.join(temp_dir, file)
                        # 永続的な場所にコピー
                        final_pdf_path = tempfile.mktemp(suffix='.pdf')
                        shutil.copy2(pdf_path, final_pdf_path)
                        shutil.rmtree(temp_dir)
                        print(f"[INFO] PDFに変換完了: {final_pdf_path}")
                        return final_pdf_path
            
            shutil.rmtree(temp_dir)
            print(f"[ERROR] PDF変換失敗: {result.stderr}")
            return None
            
        except Exception as e:
            print(f"[ERROR] PDF変換エラー: {e}")
            return None
    
    def _convert_pdf_to_png(self, pdf_path, output_path):
        """PDFをPNGに変換（PyMuPDF使用）"""
        try:
            debug_print("[DEBUG] PyMuPDFでPDF→PNG変換実行...")
            
            doc = fitz.open(pdf_path)
            if len(doc) == 0:
                print("[ERROR] PDFにページが含まれていません")
                doc.close()
                return False
            
            page = doc[0]
            
            mat = fitz.Matrix(300/72, 300/72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            img_data = pix.tobytes("png")
            pix = None
            doc.close()
            
            img = Image.open(io.BytesIO(img_data))
            
            if img.mode == 'RGBA':
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[3] if len(img.split()) > 3 else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            img = self._trim_white_margins(img)
            
            width, height = img.size
            new_width = int(width * 2)
            new_height = int(height * 2)
            img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            img.save(output_path, 'PNG', quality=95)
            
            print(f"[INFO] PNG変換完了: {output_path} (サイズ: {img.size[0]}x{img.size[1]})")
            return True
                
        except Exception as e:
            print(f"[ERROR] PNG変換エラー: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _trim_white_margins(self, img):
        """画像の白い余白をトリミング"""
        import numpy as np
        
        img_array = np.array(img)
        
        if len(img_array.shape) == 3:
            gray = np.mean(img_array, axis=2)
        else:
            gray = img_array
        
        threshold = 250
        non_white_pixels = gray < threshold
        
        rows = np.any(non_white_pixels, axis=1)
        cols = np.any(non_white_pixels, axis=0)
        
        if not rows.any() or not cols.any():
            return img
        
        row_indices = np.where(rows)[0]
        col_indices = np.where(cols)[0]
        
        top = row_indices[0]
        bottom = row_indices[-1] + 1
        left = col_indices[0]
        right = col_indices[-1] + 1
        
        return img.crop((left, top, right, bottom))
    
    def _convert_pdf_to_svg(self, pdf_path, output_path):
        """PDFをSVGに変換（PyMuPDF使用、コンテンツ領域にクロップ）"""
        try:
            import numpy as np
            from PIL import Image as PILImage
            import re
            
            debug_print("[DEBUG] PyMuPDFでPDF→SVG変換実行...")
            
            doc = fitz.open(pdf_path)
            if len(doc) == 0:
                print("[ERROR] PDFにページが含まれていません")
                doc.close()
                return False
            
            page = doc[0]
            
            # 1. 高解像度PNGとしてレンダリングしてコンテンツ領域を検出
            dpi = 300
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            
            img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img_array = np.array(img)
            
            # 2. 非白領域のbboxをピクセル単位で取得
            if len(img_array.shape) == 3:
                gray = np.mean(img_array, axis=2)
            else:
                gray = img_array
            
            threshold = 250
            non_white_pixels = gray < threshold
            
            rows = np.any(non_white_pixels, axis=1)
            cols = np.any(non_white_pixels, axis=0)
            
            # 3. SVG生成
            svg_content = page.get_svg_image()
            width_units = page.rect.width
            height_units = page.rect.height
            
            if rows.any() and cols.any():
                row_indices = np.where(rows)[0]
                col_indices = np.where(cols)[0]
                
                top_px = row_indices[0]
                bottom_px = row_indices[-1] + 1
                left_px = col_indices[0]
                right_px = col_indices[-1] + 1
                
                # 4. ピクセル座標 → viewBox座標への変換
                scale_x = width_units / pix.width
                scale_y = height_units / pix.height
                
                left_u = left_px * scale_x
                top_u = top_px * scale_y
                width_u = (right_px - left_px) * scale_x
                height_u = (bottom_px - top_px) * scale_y
                
                # 5. SVGのroot <svg> の viewBox / width / height を置き換え
                svg_content = self._update_svg_viewbox(
                    svg_content, left_u, top_u, width_u, height_u
                )
            
            doc.close()
            
            # SVGファイルに書き込み
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(svg_content)
            
            print(f"[INFO] SVG変換完了: {output_path}")
            return True
                
        except Exception as e:
            print(f"[ERROR] SVG変換エラー: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _update_svg_viewbox(self, svg_content, left, top, width, height, scale=2.0):
        """SVGのviewBoxとwidth/heightを更新する
        
        Args:
            svg_content: SVG文字列
            left, top, width, height: 新しいviewBox座標
            scale: 表示サイズの倍率（デフォルト: 2.0、PNGと同じ）
            
        Returns:
            str: 更新されたSVG文字列
        """
        import re
        
        # viewBox属性を更新（座標系はそのまま）
        new_viewbox = f'viewBox="{left:.2f} {top:.2f} {width:.2f} {height:.2f}"'
        svg_content = re.sub(
            r'viewBox="[^"]*"',
            new_viewbox,
            svg_content,
            count=1
        )
        
        # width属性を更新（表示サイズをscale倍に拡大）
        display_width = width * scale
        svg_content = re.sub(
            r'width="[^"]*"',
            f'width="{display_width:.2f}"',
            svg_content,
            count=1
        )
        
        # height属性を更新（表示サイズをscale倍に拡大）
        display_height = height * scale
        svg_content = re.sub(
            r'height="[^"]*"',
            f'height="{display_height:.2f}"',
            svg_content,
            count=1
        )
        
        return svg_content
    
    def _convert_pdf_to_image(self, pdf_path, output_path):
        """PDFを画像に変換（出力形式に応じてPNGまたはSVG）"""
        if self.output_format == 'svg':
            # SVG出力の場合は拡張子を変更
            svg_path = output_path.replace('.png', '.svg')
            return self._convert_pdf_to_svg(pdf_path, svg_path), svg_path
        else:
            return self._convert_pdf_to_png(pdf_path, output_path), output_path
    
    def _debug_pdf_content(self, pdf_path):
        """PDFの内容をデバッグ（オプション）"""
        try:
            # pdfinfoが利用可能な場合のみ実行
            if shutil.which('pdfinfo'):
                cmd = ['pdfinfo', pdf_path]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                if result.returncode == 0:
                    # ファイルサイズなど基本情報のみ表示
                    lines = result.stdout.split('\n')
                    for line in lines:
                        if 'Pages:' in line or 'Page size:' in line:
                            debug_print(f"[DEBUG] {line.strip()}")
                else:
                    debug_print(f"[DEBUG] PDF情報取得失敗")
            # pdfinfoがない場合は何もしない（エラーメッセージなし）
        except Exception:
            # エラーが発生しても無視（オプション機能のため）
            pass
    
    def _debug_image_info(self, image_path):
        """生成された画像の詳細情報をデバッグ"""
        try:
            from PIL import Image
            with Image.open(image_path) as img:
                info_parts = [
                    f"サイズ: {img.size[0]}x{img.size[1]}",
                    f"モード: {img.mode}",
                ]
                if hasattr(img, 'info') and 'dpi' in img.info:
                    info_parts.append(f"DPI: {img.info['dpi']}")
                debug_print(f"[DEBUG] 画像情報: {' | '.join(info_parts)}")
        except Exception as e:
            debug_print(f"[DEBUG] 画像情報取得エラー: {e}")

    def _extract_shape_metadata_from_drawing(self, drawing_element) -> Dict[str, Any]:
        """DrawingML要素から図形メタデータを抽出"""
        metadata = {
            'type': 'unknown',
            'name': '',
            'description': '',
            'shapes': []
        }
        
        try:
            processed_ids = set()
            
            for elem in drawing_element.iter():
                tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                
                if tag_name == 'wgp':
                    for child in elem:
                        child_tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                        
                        if child_tag == 'wsp':
                            shape_info = self._extract_wsp_metadata(child)
                            if shape_info and (shape_info.get('name') or shape_info.get('text') or shape_info.get('shape_type')):
                                shape_id = shape_info.get('id', '')
                                if not shape_id or shape_id not in processed_ids:
                                    metadata['shapes'].append(shape_info)
                                    if shape_id:
                                        processed_ids.add(shape_id)
                        
                        elif child_tag == 'pic':
                            shape_info = self._extract_pic_metadata(child)
                            if shape_info and (shape_info.get('name') or shape_info.get('type')):
                                shape_id = shape_info.get('id', '')
                                if not shape_id or shape_id not in processed_ids:
                                    metadata['shapes'].append(shape_info)
                                    if shape_id:
                                        processed_ids.add(shape_id)
                        
                        elif child_tag == 'grpSp':
                            for nested in child.iter():
                                nested_tag = nested.tag.split('}')[-1] if '}' in nested.tag else nested.tag
                                
                                if nested_tag == 'wsp':
                                    shape_info = self._extract_wsp_metadata(nested)
                                    if shape_info and (shape_info.get('name') or shape_info.get('text') or shape_info.get('shape_type')):
                                        shape_id = shape_info.get('id', '')
                                        if not shape_id or shape_id not in processed_ids:
                                            metadata['shapes'].append(shape_info)
                                            if shape_id:
                                                processed_ids.add(shape_id)
                                
                                elif nested_tag == 'pic':
                                    shape_info = self._extract_pic_metadata(nested)
                                    if shape_info and (shape_info.get('name') or shape_info.get('type')):
                                        shape_id = shape_info.get('id', '')
                                        if not shape_id or shape_id not in processed_ids:
                                            metadata['shapes'].append(shape_info)
                                            if shape_id:
                                                processed_ids.add(shape_id)
                
                elif tag_name == 'anchor' or tag_name == 'inline':
                    shape_info = self._extract_single_shape_metadata(elem)
                    if shape_info and shape_info.get('name'):
                        shape_id = shape_info.get('id', '')
                        if not shape_id or shape_id not in processed_ids:
                            metadata['shapes'].append(shape_info)
                            if shape_id:
                                processed_ids.add(shape_id)
            
        except Exception as e:
            debug_print(f"[DEBUG] 図形メタデータ抽出エラー: {e}")
            import traceback
            traceback.print_exc()
        
        return metadata
    
    def _extract_single_shape_metadata(self, anchor_element) -> Dict[str, Any]:
        """単一図形のメタデータを抽出"""
        shape_info = {}
        
        try:
            for elem in anchor_element.iter():
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                
                if tag == 'cNvPr':
                    shape_info['name'] = elem.attrib.get('name', '')
                    shape_info['id'] = elem.attrib.get('id', '')
                    shape_info['description'] = elem.attrib.get('descr', '')
                
                elif tag in ('rect', 'roundRect', 'ellipse', 'triangle', 'line', 'bentConnector2', 
                            'bentConnector3', 'bentConnector4', 'bentConnector5', 'straightConnector1'):
                    shape_info['shape_type'] = tag
                
                elif tag == 'txBody' or tag == 'sp':
                    text_parts = []
                    for t_elem in elem.iter():
                        t_tag = t_elem.tag.split('}')[-1] if '}' in t_elem.tag else t_elem.tag
                        if t_tag == 't' and t_elem.text:
                            text_parts.append(t_elem.text.strip())
                    if text_parts:
                        shape_info['text'] = ' / '.join(text_parts)
                
                elif tag == 'extent':
                    try:
                        cx = int(elem.attrib.get('cx', 0))
                        cy = int(elem.attrib.get('cy', 0))
                        shape_info['width_emu'] = cx
                        shape_info['height_emu'] = cy
                    except:
                        pass
                
                elif tag == 'off' or tag == 'pos':
                    try:
                        x = int(elem.attrib.get('x', 0))
                        y = int(elem.attrib.get('y', 0))
                        shape_info['x_emu'] = x
                        shape_info['y_emu'] = y
                    except:
                        pass
        
        except Exception as e:
            debug_print(f"[DEBUG] 単一図形メタデータ抽出エラー: {e}")
        
        return shape_info
    
    def _extract_wsp_metadata(self, wsp_element) -> Dict[str, Any]:
        """wsp（Word Shape）要素からメタデータを抽出"""
        shape_info = {}
        text_parts = []
        
        try:
            for elem in wsp_element.iter():
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                
                if tag == 'cNvPr':
                    shape_info['name'] = elem.attrib.get('name', '')
                    shape_info['id'] = elem.attrib.get('id', '')
                    shape_info['description'] = elem.attrib.get('descr', '')
                
                elif tag == 'prstGeom':
                    prst = elem.attrib.get('prst', '')
                    if prst:
                        shape_info['shape_type'] = prst
                
                elif tag == 't' and elem.text:
                    text_parts.append(elem.text.strip())
                
                elif tag == 'ext':
                    try:
                        cx = int(elem.attrib.get('cx', 0))
                        cy = int(elem.attrib.get('cy', 0))
                        shape_info['width_emu'] = cx
                        shape_info['height_emu'] = cy
                    except:
                        pass
                
                elif tag == 'off':
                    try:
                        x = int(elem.attrib.get('x', 0))
                        y = int(elem.attrib.get('y', 0))
                        shape_info['x_emu'] = x
                        shape_info['y_emu'] = y
                    except:
                        pass
            
            if text_parts:
                shape_info['text'] = ' / '.join(text_parts)
        
        except Exception as e:
            debug_print(f"[DEBUG] wsp要素メタデータ抽出エラー: {e}")
        
        return shape_info
    
    def _extract_pic_metadata(self, pic_element) -> Dict[str, Any]:
        """pic（Picture）要素からメタデータを抽出"""
        shape_info = {}
        
        try:
            for elem in pic_element.iter():
                tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
                
                if tag == 'cNvPr':
                    shape_info['name'] = elem.attrib.get('name', '')
                    shape_info['id'] = elem.attrib.get('id', '')
                    shape_info['description'] = elem.attrib.get('descr', '')
                    shape_info['type'] = 'picture'
                
                elif tag == 'ext':
                    try:
                        cx = int(elem.attrib.get('cx', 0))
                        cy = int(elem.attrib.get('cy', 0))
                        shape_info['width_emu'] = cx
                        shape_info['height_emu'] = cy
                    except:
                        pass
                
                elif tag == 'off':
                    try:
                        x = int(elem.attrib.get('x', 0))
                        y = int(elem.attrib.get('y', 0))
                        shape_info['x_emu'] = x
                        shape_info['y_emu'] = y
                    except:
                        pass
                
                elif tag == 'blip':
                    embed = elem.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed', '')
                    if embed:
                        shape_info['image_rel_id'] = embed
        
        except Exception as e:
            debug_print(f"[DEBUG] pic要素メタデータ抽出エラー: {e}")
        
        return shape_info
    
    def _format_shape_metadata_as_text(self, metadata: Dict[str, Any]) -> str:
        """図形メタデータを人間が読みやすいテキスト形式に整形"""
        if not metadata.get('shapes'):
            return ""
        
        lines = ["### 図形情報", ""]
        
        for idx, shape in enumerate(metadata['shapes'], 1):
            name = shape.get('name', '')
            if not name:
                shape_type = shape.get('shape_type', shape.get('type', ''))
                if shape_type:
                    name = f"図形 #{idx} ({shape_type})"
                else:
                    name = f"図形 #{idx}"
            
            lines.append(f"**{name}**")
            
            if shape.get('id'):
                lines.append(f"- ID: {shape['id']}")
            
            if shape.get('shape_type'):
                lines.append(f"- 図形タイプ: {shape['shape_type']}")
            elif shape.get('type'):
                lines.append(f"- タイプ: {shape['type']}")
            
            if shape.get('text'):
                lines.append(f"- テキスト: {shape['text']}")
            
            if shape.get('description'):
                lines.append(f"- 説明: {shape['description']}")
            
            lines.append("")
        
        return '\n'.join(lines)
    
    def _format_shape_metadata_as_json(self, metadata: Dict[str, Any]) -> str:
        """図形メタデータをJSON形式に整形"""
        import json
        if not metadata.get('shapes'):
            return "{}"
        return json.dumps(metadata, ensure_ascii=False, indent=2)
    
    def _detect_image_format(self, image_data: bytes, target_ref: str) -> str:
        """画像形式を検出"""
        if image_data.startswith(b'\x89PNG'):
            return '.png'
        elif image_data.startswith(b'\xff\xd8\xff'):
            return '.jpg'
        elif image_data.startswith(b'GIF'):
            return '.gif'
        elif image_data.startswith(b'\x01\x00\x00\x00'):
            return '.emf'
        elif image_data.startswith(b'\xd7\xcd\xc6\x9a'):
            return '.wmf'
        
        # target_refから推測
        if target_ref.endswith('.emf'):
            return '.emf'
        elif target_ref.endswith('.wmf'):
            return '.wmf'
        elif target_ref.endswith('.jpeg'):
            return '.jpg'
        
        return '.png'  # デフォルト
    
    def _convert_vector_image(self, image_data: bytes, original_path: str) -> Optional[str]:
        """ベクター画像を画像に変換（出力形式に応じてPNGまたはSVG）"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(original_path).suffix) as temp_file:
                temp_file.write(image_data)
                temp_path = temp_file.name
            
            # 出力形式に応じて拡張子を決定
            fmt = getattr(self, 'output_format', 'png')
            ext = '.svg' if fmt == 'svg' else '.png'
            output_path = original_path.replace('.emf', ext).replace('.wmf', ext)
            
            if self._convert_with_libreoffice(temp_path, output_path):
                os.unlink(temp_path)
                return output_path
            
            os.unlink(temp_path)
            print(f"[ERROR] ベクター画像変換失敗")
            with open(original_path, 'wb') as f:
                f.write(image_data)
            return original_path
                
        except Exception as e:
            print(f"[ERROR] ベクター画像変換エラー: {e}")
            with open(original_path, 'wb') as f:
                f.write(image_data)
            return original_path

    def _convert_with_libreoffice(self, input_path: str, output_path: str) -> bool:
        """LibreOfficeを使用してベクター画像を変換（出力形式に応じてPNGまたはSVG）"""
        temp_dir = None
        try:
            temp_dir = tempfile.mkdtemp()
            
            cmd = [
                LIBREOFFICE_PATH,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', temp_dir,
                input_path
            ]
            
            subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            pdf_path = None
            for file in os.listdir(temp_dir):
                if file.endswith('.pdf'):
                    pdf_path = os.path.join(temp_dir, file)
                    break
            
            if pdf_path and os.path.exists(pdf_path):
                # 出力形式を判定（拡張子から）
                fmt = 'svg' if output_path.endswith('.svg') else 'png'
                
                if fmt == 'svg':
                    # SVG出力
                    if self._convert_pdf_to_svg(pdf_path, output_path):
                        print(f"[SUCCESS] ベクター画像変換完了（LibreOffice→PDF→SVG）: {output_path}")
                        return True
                else:
                    # PNG出力（既存の処理）
                    pdf_doc = fitz.open(pdf_path)
                    page = pdf_doc[0]
                    
                    mat = fitz.Matrix(300 / 72, 300 / 72)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    
                    from PIL import Image
                    if pix.alpha:
                        img = Image.frombytes("RGBA", [pix.width, pix.height], pix.samples)
                        bg = Image.new("RGB", img.size, (255, 255, 255))
                        bg.paste(img, mask=img.split()[3])
                    else:
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    pdf_doc.close()
                    
                    img = self._trim_white_margins(img)
                    img.save(output_path, "PNG")
                    
                    if os.path.exists(output_path):
                        print(f"[SUCCESS] ベクター画像変換完了（LibreOffice→PDF→PNG）: {output_path}")
                        return True
            
            return False
            
        except Exception as e:
            print(f"[ERROR] LibreOffice変換エラー: {e}")
            return False
        finally:
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)


def convert_doc_to_docx(doc_file_path: str) -> str:
    """DOCファイルをDOCXに変換"""
    import subprocess
    import tempfile
    import shutil
    from pathlib import Path
    
    print(f"[INFO] DOCファイルをDOCXに変換中: {doc_file_path}")
    
    # 一時ディレクトリを作成（自動削除されない）
    temp_dir = tempfile.mkdtemp(prefix='word2md_doc_conversion_')
    
    try:
        # 出力ファイル名を決定
        doc_path = Path(doc_file_path)
        docx_filename = doc_path.stem + '.docx'
        docx_output_path = os.path.join(temp_dir, docx_filename)
        
        # LibreOfficeを使用してDOCをDOCXに変換
        cmd = [
            LIBREOFFICE_PATH,
            '--headless',
            '--convert-to', 'docx',
            '--outdir', temp_dir,
            doc_file_path
        ]
        
        debug_print(f"[DEBUG] LibreOffice変換コマンド: {' '.join(cmd)}")
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode != 0:
            print(f"[ERROR] LibreOffice変換失敗: {result.stderr}")
            shutil.rmtree(temp_dir)  # エラー時は一時ディレクトリを削除
            return None
        
        # 変換されたファイルが存在するか確認
        if not os.path.exists(docx_output_path):
            print(f"[ERROR] 変換後のDOCXファイルが見つかりません: {docx_output_path}")
            shutil.rmtree(temp_dir)  # エラー時は一時ディレクトリを削除
            return None
        
        print(f"[SUCCESS] DOC→DOCX変換完了: {docx_output_path}")
        print(f"[INFO] 一時ファイル作成: {docx_output_path}")
        
        # 一時ディレクトリのパスを返す（後でメイン関数で削除）
        return docx_output_path
        
    except subprocess.TimeoutExpired:
        print("[ERROR] LibreOffice変換がタイムアウトしました")
        shutil.rmtree(temp_dir)  # エラー時は一時ディレクトリを削除
        return None
    except Exception as e:
        print(f"[ERROR] DOC変換エラー: {e}")
        shutil.rmtree(temp_dir)  # エラー時は一時ディレクトリを削除
        return None


def main():
    """メイン関数"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Word文書をMarkdownに変換')
    parser.add_argument('word_file', help='変換するWord文書ファイル（.docx/.doc）')
    parser.add_argument('--use-heading-text', action='store_true', 
                       help='章番号の代わりに見出しテキストをリンクに使用')
    parser.add_argument('-o', '--output-dir', type=str, 
                       help='出力ディレクトリを指定（デフォルト: 実行ディレクトリ）')
    parser.add_argument('--shape-metadata', action='store_true',
                       help='図形メタデータを画像の後に出力（テキスト形式とJSON形式）')
    parser.add_argument('--format', choices=['png', 'svg'], default='svg',
                       help='出力画像形式を指定（デフォルト: png）')
    parser.add_argument('-v', '--verbose', action='store_true',
                       help='デバッグ情報を出力')
    
    args = parser.parse_args()
    
    set_verbose(args.verbose)
    
    if not os.path.exists(args.word_file):
        print(f"エラー: ファイル '{args.word_file}' が見つかりません。")
        sys.exit(1)
    
    if not args.word_file.endswith(('.docx', '.doc')):
        print("エラー: .docxまたは.doc形式のファイルを指定してください。")
        sys.exit(1)
    
    # DOCファイルの場合は事前にDOCXに変換
    processing_file = args.word_file
    converted_file = None
    
    if args.word_file.endswith('.doc'):
        print("DOCファイルが指定されました。DOCXに変換します...")
        converted_file = convert_doc_to_docx(args.word_file)
        if converted_file is None:
            print("DOC→DOCX変換に失敗しました。")
            sys.exit(1)
        processing_file = converted_file
        print(f"DOC→DOCX変換完了: {converted_file}")
    
    try:
        converter = WordToMarkdownConverter(
            processing_file, 
            use_heading_text=args.use_heading_text, 
            output_dir=args.output_dir, 
            shape_metadata=args.shape_metadata,
            output_format=args.format
        )
        output_file = converter.convert()
        print("\n変換完了!")
        print(f"出力ファイル: {output_file}")
        print(f"画像フォルダ: {converter.images_dir}")
        if args.use_heading_text:
            print("見出しテキストリンクモード: 有効")
        
    except Exception as e:
        print(f"変換エラー: {e}")
        sys.exit(1)
    finally:
        # 一時的に作成したDOCXファイルとその親ディレクトリを必ず削除
        if converted_file:
            try:
                import shutil
                from pathlib import Path
                
                # 一時ディレクトリのパスを取得
                temp_dir = Path(converted_file).parent
                
                if temp_dir.exists() and temp_dir.name.startswith('word2md_doc_conversion_'):
                    shutil.rmtree(temp_dir)
                    print(f"一時ディレクトリを削除: {temp_dir}")
                elif os.path.exists(converted_file):
                    os.remove(converted_file)
                    print(f"一時ファイルを削除: {converted_file}")
                    
            except Exception as cleanup_error:
                print(f"一時ファイル削除に失敗: {cleanup_error}")


if __name__ == "__main__":
    main()

